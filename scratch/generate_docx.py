import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_resume_docx():
    doc = Document()
    
    # Page setup - 0.75" margins (narrow/standard for dense resumes)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Color Palette: Slate / Corporate (Deep Navy accents)
    COLOR_PRIMARY = RGBColor(12, 74, 110)   # #0C4A6E (Deep Navy for headings)
    COLOR_BODY = RGBColor(51, 65, 85)       # #334155 (Slate Ink for text)
    COLOR_MUTED = RGBColor(100, 116, 139)   # #64748B (Muted gray for details)

    # Set default paragraph format
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = COLOR_BODY

    # Helper function to add headings with styling & bottom border
    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_PRIMARY
        
        # Add a subtle horizontal line under section heading
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                         r'<w:bottom w:val="single" w:sz="6" w:space="4" w:color="CBD5E1"/>'
                         r'</w:pBdr>')
        p._p.get_or_add_pPr().append(pBdr)
        return p

    # Read Markdown resume
    md_path = r"E:\1_Resume_Prepreation\OLD_Resume\Chokkar_Gurusamy_Resume.md"
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_summary = False
    in_skills = False

    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue
            
        # Header processing
        if line_str.startswith("# "):
            name = line_str[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = COLOR_PRIMARY
            continue
            
        if "mailto:" in line_str or "linkedin.com" in line_str:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            
            # Simple clean link text extraction
            parts = [part.strip() for part in line_str.split('|')]
            for i, part in enumerate(parts):
                clean_part = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', part) # strip markdown links
                run = p.add_run(clean_part)
                run.font.size = Pt(9.5)
                run.font.color.rgb = COLOR_MUTED
                if i < len(parts) - 1:
                    p.add_run("   |   ")
            continue

        if line_str == "---":
            continue

        # Headings
        if line_str.startswith("## "):
            title = line_str[3:]
            add_section_heading(title)
            continue

        # Sub-headings (Companies / Projects / Degree)
        if line_str.startswith("### "):
            sub_title = line_str[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            
            # Formats "Name | Title | Date" with distinct bolding
            parts = [part.strip() for part in sub_title.split('|')]
            if len(parts) >= 2:
                # Bold Company / Project Name
                r1 = p.add_run(parts[0])
                r1.bold = True
                
                # Add Middle Role
                r2 = p.add_run(f" | {parts[1]}")
                r2.italic = True
                
                # Add Date if present
                if len(parts) == 3:
                    p.add_run(" | ")
                    r3 = p.add_run(parts[2])
                    r3.font.color.rgb = COLOR_MUTED
            else:
                # Single part subhead
                r = p.add_run(sub_title)
                r.bold = True
            continue

        # List items / bullet points
        if line_str.startswith("- ") or line_str.startswith("* "):
            bullet_text = line_str[2:]
            
            # Clean up inline bold formatting (**word**)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            
            parts = re.split(r'(\*\*.*?\*\*)', bullet_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    clean_text = part[2:-2]
                    r = p.add_run(clean_text)
                    r.bold = True
                else:
                    p.add_run(part)
            continue

        # Standard text paragraphs / Summaries
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        
        # Parse inline bold (**bold**)
        parts = re.split(r'(\*\*.*?\*\*)', line_str)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                clean_text = part[2:-2]
                r = p.add_run(clean_text)
                r.bold = True
            else:
                p.add_run(part)

    # Save output to same directory
    output_path = r"E:\1_Resume_Prepreation\OLD_Resume\Chokkar_Gurusamy_Resume.docx"
    doc.save(output_path)
    print(f"Successfully generated docx at: {output_path}")

if __name__ == "__main__":
    create_resume_docx()
