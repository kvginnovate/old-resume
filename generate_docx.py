import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_resume():
    doc = Document()
    
    # 1. Page Margins (Tight/Compact to fit 2 pages nicely, similar to LaTeX)
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    # Styles & Colors
    COLOR_PRIMARY = RGBColor(15, 23, 42)      # Slate 900
    COLOR_MUTED = RGBColor(71, 85, 105)       # Slate 600
    COLOR_ACCENT = RGBColor(37, 99, 235)      # Royal Blue
    
    # Configure default style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = COLOR_PRIMARY
    
    # Helper to add bottom border to paragraph (horizontal rule)
    def add_bottom_border(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="4" w:color="E2E8F0"/></w:pBdr>')
        pPr.append(pBdr)

    # Helper to add section headings
    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_PRIMARY
        add_bottom_border(p)
        return p

    # Helper for alignment tables (Left title, Right date/tech)
    def add_heading_row(left_text, right_text, is_bold=True, is_italic=False, space_before=4):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        # Total width = 8.5 - 1.3 = 7.2 inches
        table.columns[0].width = Inches(5.2)
        table.columns[1].width = Inches(2.0)
        
        row = table.rows[0]
        # Remove cell padding
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for m in ['top', 'bottom', 'left', 'right']:
                node = OxmlElement(f'w:{m}')
                node.set(qn('w:w'), '0')
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)

        cell_l, cell_r = row.cells[0], row.cells[1]
        
        p_l = cell_l.paragraphs[0]
        p_l.paragraph_format.space_before = Pt(space_before)
        p_l.paragraph_format.space_after = Pt(2)
        
        p_r = cell_r.paragraphs[0]
        p_r.paragraph_format.space_before = Pt(space_before)
        p_r.paragraph_format.space_after = Pt(2)
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run_l = p_l.add_run(left_text)
        run_l.bold = is_bold
        run_l.font.italic = is_italic
        
        run_r = p_r.add_run(right_text)
        run_r.bold = is_bold
        run_r.font.italic = is_italic
        
        return table

    # Helper to add bullet point with inline formatting
    def add_bullet(runs_data):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing = 1.15
        
        for text, bold, italic in runs_data:
            run = p.add_run(text)
            if bold:
                run.bold = True
            if italic:
                run.font.italic = True
        return p

    # --- HEADER ---
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run("CHOKKAR GURUSAMY")
    run_name.bold = True
    run_name.font.size = Pt(22)
    run_name.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(2)
    run_sub = p_sub.add_run("Bangalore, India")
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = COLOR_MUTED

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(8)
    p_contact.paragraph_format.line_spacing = 1.0
    
    contacts = [
        ("+91 99408 92858", None),
        ("chokkar.g@gmail.com", "mailto:chokkar.g@gmail.com"),
        ("linkedin.com/in/chokkarg", "https://linkedin.com/in/chokkarg"),
        
    ]
    for i, (text, url) in enumerate(contacts):
        if i > 0:
            p_contact.add_run("   |   ").font.color.rgb = COLOR_MUTED
        run = p_contact.add_run(text)
        run.font.size = Pt(9.5)
        if url:
            run.font.color.rgb = COLOR_ACCENT
            run.underline = True

    # --- PROFESSIONAL SUMMARY ---
    add_section_heading("Professional Summary")
    p_summary = doc.add_paragraph()
    p_summary.paragraph_format.space_after = Pt(6)
    p_summary.paragraph_format.line_spacing = 1.15
    
    summary_parts = [
        ("Staff Engineer with ", False, False),
        ("14+ years of experience", True, False),
        (" owning end-to-end solution architecture for enterprise platforms, including API services, event-driven integration, data access, security, observability, and operational readiness. At Dish Network, own architecture for the subscriber-facing microservices platform serving ", False, False),
        ("10M+ subscribers", True, False),
        (" on Rancher-managed Kubernetes. Promoted from ", False, False),
        ("Lead to Staff (Solution Architect)", True, False),
        (" for org-wide Spring Boot 3 migration, Snyk security standards, and engineering practices adopted across teams. Mentor ", False, False),
        ("engineers", True, False),
        (" and lead architecture reviews, incident response, and senior/staff hiring. Hands-on in Java/Spring Boot REST APIs, OpenAPI, Kafka/event-driven patterns, cloud-native delivery, and quality gates, AI-augmented SDLC, and AI governance practices. Regularly leverage AWS Kiro for architecture analysis, automated documentation/flow-diagram generation, and subagent-based code workflows with cross-session AI memory persistence. Introduced AI-powered Spring Boot modernization to Dish leadership by presenting a live demo of Kiro spec-driven development, automated code generation, and subagent orchestration to VP+ engineering leadership.", False, False)
    ]
    for text, bold, italic in summary_parts:
        run = p_summary.add_run(text)
        run.bold = bold
        run.font.italic = italic
        run.font.size = Pt(10)

    # --- TECHNICAL SKILLS ---
    add_section_heading("Technical Skills")
    skills = [
        ("Languages", "Java, Kotlin, JavaScript, TypeScript, Python, SQL"),
        ("Backend & APIs", "Spring Boot 3, Spring Cloud, Hibernate, REST APIs, OpenAPI/Swagger, Resilience4j, Liquibase, Maven"),
        ("Architecture & Integration", "Microservices, Event-Driven Architecture, Kafka, RabbitMQ, API Gateway, Circuit Breaker, DDD, ADRs"),
        ("Frontend & Mobile", "ReactJS, Next.js, React Native, Android SDK, and API Integration Patterns"),
        ("Data", "PostgreSQL, SQL, Oracle, MongoDB, DynamoDB, GemFire (distributed cache), Firebase"),
        ("Cloud & Platform", "Docker, Kubernetes (Rancher), GCP, AWS, and Multi-Cloud Architecture (Azure/AKS)"),
        ("DevOps, Security & Quality", "GitLab CI/CD, Jenkins, Snyk, Dynatrace, Prometheus, ELK, JUnit, Selenium, Postman, API & performance testing"),
        ("AI & Automation", "Kiro (spec-driven dev, hooks, subagents, EARS), MCP integrations (GitLab, Jira, GitHub), custom agent skills (defect-fix / test-gen / code-review), LLM workflows, AI-augmented SDLC (AI code review, automated test generation, spec-to-code traceability, AI observability)")
    ]
    for category, items in skills:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.left_indent = Inches(0.15)
        
        run_cat = p.add_run(f"{category}: ")
        run_cat.bold = True
        run_cat.font.size = Pt(9.5)
        
        run_items = p.add_run(items)
        run_items.font.size = Pt(9.5)

    # --- PROFESSIONAL EXPERIENCE ---
    add_section_heading("Professional Experience")
    
    # 1. Staff Engineer
    add_heading_row("Dish Network Technologies | Staff Engineer (promoted from Lead)", "May 2022 \u2013 Present", is_bold=True, space_before=6)
    
    add_bullet([
        ("Winner of the ", False, False),
        ("Annual Technical Excellence Awards", True, False),
        (" (2024, 2025) \u00b7 Recipient of the ", False, False),
        ("Spot Awards", True, False),
        (" (2023, 2025) \u00b7 ", False, False),
        ("CPAW Award", True, False),
        (" recipient, recognizing contributions to delivery impact.", False, False)
    ])

    # 2. My Dish App
    add_heading_row("My Dish App \u2013 Subscriber Platform", "", space_before=4)
    add_bullet([
        ("Architected enterprise REST microservices (API-first / OpenAPI) for the flagship subscriber platform serving ", False, False),
        ("10M+ subscribers", True, False),
        (" to maintain ", False, False),
        ("99.9% uptime", True, False),
        (" with sub-second API latency, and defined integration patterns across frontend, backend, and shared platform services.", False, False)
    ])
    add_bullet([
        ("Drove the Spring Boot 3 migration across platform services using Kiro spec-driven modernization, automating the remediation of breaking changes across ", False, False),
        ("50K+ lines of code", True, False),
        (" to achieve a ", False, False),
        ("60% productivity gain", True, False),
        (", and built a custom dashboard to track AI-committed code changes.", False, False)
    ])
    add_bullet([
        ("Leveraged spec-driven AI workflows to accelerate backend feature development, delivering high-quality microservices with automated spec-to-code traceability.", False, False)
    ])
    add_bullet([
        ("Established security and operational gates by integrating a custom ", False, False),
        ("Dynatrace MCP server", True, False),
        (" for automated incident analysis, distributed tracing, and a ", False, False),
        ("Snyk MCP server", True, False),
        (" paired with ", False, False),
        ("Kiro agents", True, False),
        (" to scan and remediate vulnerabilities (achieving same-day fixes), while mandating ", False, False),
        ("85%+", True, False),
        (" JUnit/Postman coverage.", False, False)
    ])

    # 3. Sling TV
    add_heading_row("Sling TV Product Development", "", space_before=4)
    add_bullet([
        ("Led the strategic transformation of ", False, False),
        ("80+ legacy Ruby on Rails APIs", True, False),
        (" into a modern, scalable Java microservices architecture, achieving a ", False, False),
        ("50% faster delivery timeline", True, False),
        (" through API-first patterns and event-driven design.", False, False)
    ])
    add_bullet([
        ("Orchestrated end-to-end integration of ", False, False),
        ("Apple Pay", True, False),
        (" and ", False, False),
        ("Google Pay", True, False),
        (" payment ecosystems, driving direct organizational revenue growth.", False, False)
    ])
    add_bullet([
        ("Spearheaded GenAI modernization by leveraging Amazon Q to reverse-engineer ", False, False),
        ("400K+ lines", True, False),
        (" of legacy code and scaffold ", False, False),
        ("2M+ lines", True, False),
        (" of Java microservices and ", False, False),
        ("10K+ unit tests", True, False),
        (" (maintaining ", False, False),
        ("85% coverage", True, False),
        ("), achieving a ", False, False),
        ("70% productivity gain", True, False),
        (".", False, False)
    ])
    add_bullet([
        ("Designed and implemented resilient Kafka event-driven architectures for payment and billing events, ensuring zero-loss message delivery across subscriber services.", False, False)
    ])

    # 4. Dish Paperless Agreement (DPA), STB Health, Field Catalogue & App Store
    add_heading_row("Dish Paperless Agreement (DPA), STB Health, Field Catalogue & App Store", "", space_before=4)
    add_bullet([
        ("Dish Paperless Agreement (DPA):", True, False),
        (" Led the migration of the paperless onboarding system from ColdFusion to Spring Boot microservices, enabling field agents to sign agreements, cross-sell services, and activate equipment offline while orchestrating agreements through 23 integrations and 7 SFTP distribution targets.", False, False)
    ])
    add_bullet([
        ("STB Health Monitor:", True, False),
        (" Real-time dashboards and REST APIs for Set-Top-Box telemetry across ", False, False),
        ("10M+ devices", True, False),
        (", enabling proactive ops before subscriber impact.", False, False)
    ])
    add_bullet([
        ("Field Catalogue:", True, False),
        (" Offline-first Android app with ", False, False),
        ("100%", True, False),
        (" offline capability, enabling catalog access in remote areas with zero connectivity and contributing to a ", False, False),
        ("60%", True, False),
        (" increase in on-site technical sales; earned appreciation from the ", False, False),
        ("Sales VP", True, False),
        (", while the integrated CMS sync cut catalog update delivery time from weeks to minutes.", False, False)
    ])
    add_bullet([
        ("DISH App Store:", True, False),
        (" Custom-built internal enterprise app distribution platform (Lambda + DynamoDB + React) that delivers iOS/Android builds to DISH employees, technicians, and retailers via OTA installation with Okta SSO and group-based access control.", False, False)
    ])

    # 5. Asset Management & Shift Allowance -- Internal Tools
    add_heading_row("Asset Management & Shift Allowance -- Internal Tools", "", space_before=4)
    add_bullet([
        ("Asset Management (Won Codefest 2023):", True, False),
        (" Architected and built a full-stack system (Spring Boot, ReactJS) during a hackathon, containerizing it (Docker, Portainer) to deploy on internal VMs via JFrog, which automated tracking to cut manual effort by ", False, False),
        ("70%", True, False),
        (" and achieve company-wide adoption.", False, False)
    ])
    add_bullet([
        ("Shift Allowance:", True, False),
        (" Developed a full-stack portal (Spring Boot, ReactJS) that automated long-term manual Excel workflows to secure an HR ", False, False),
        ("Spot Award", True, False),
        (", now used company-wide by all DISH employees for shift allowance requests and payroll approvals.", False, False)
    ])

    # 6. Kiro Mobile & Custom Agent Skill Development (Dish Ignite 2026 Showcase)
    add_heading_row("Kiro Mobile & Custom Agent Skill Development (Dish Ignite 2026 Showcase)", "", space_before=4)
    add_bullet([
        ("Developed a lightweight mobile interface to monitor and control Kiro IDE agent sessions from a phone over LAN, featuring a live preview of chat, tasks, and code via Chrome DevTools Protocol, which was showcased at ", False, False),
        ("Dish Ignite 2026", True, False),
        (".", False, False)
    ])
    add_bullet([
        ("Created custom agent skills to manage and orchestrate MCP integrations (GitLab, Jira, GitHub) to automate end-to-end software delivery, including defect resolution, test generation, and automated PR/MR creation.", False, False)
    ])

    # 7. MSys
    add_heading_row("MSys Technologies (now Aziro) | Technical Architect / Lead", "Feb 2012 \u2013 May 2022", is_bold=True, space_before=6)
    add_bullet([
        ("Progressed from ", False, False),
        ("Software Engineer to Senior, Tech Lead, and Technical Architect", True, False),
        (" over 10 years serving global enterprise clients (Pivot3, Nasuni, NetApp) in storage, virtualization, and cloud.", False, False)
    ])
    add_bullet([
        ("Led engineering teams while owning scoping, architecture, demos, and delivery, successfully expanding client engagements.", False, False)
    ])
    add_bullet([
        ("Best Performer of the Year", True, False),
        (" (5x: 2013, 2014, 2016, 2018, 2019) \u00b7 ", False, False),
        ("Blog Championship Award", True, False),
        (" (2017)", False, False)
    ])

    # 8. Pivot3
    add_heading_row("Pivot3 \u2013 VMware HCI Platform", "", space_before=4)
    add_bullet([
        ("Served as Technical Lead and Full-Stack Developer for VMware integration, delivering backend REST APIs, a ReactJS vSphere UI plugin, and vRealize Orchestrator automation to achieve a ", False, False),
        ("60%", True, False),
        (" reduction in manual provisioning effort for ", False, False),
        ("1000+ volumes", True, False),
        (".", False, False)
    ])
    add_bullet([
        ("Partnered with global product stakeholders on roadmap, architecture reviews, and release delivery, translating requirements into scalable technical designs.", False, False)
    ])

    # 9. Enterprise Mobile
    add_heading_row("Enterprise Mobile & Test Automation", "", space_before=4)
    add_bullet([
        ("Led the development and launch of Nasuni and Spree Wearables Android apps (Play Store) with Firebase sync and analytics for crash/engagement monitoring.", False, False)
    ])
    add_bullet([
        ("Built Mobitaz mobile test automation (record/playback) for client demos, and delivered a NetApp Selenium regression suite covering ", False, False),
        ("100+ scenarios", True, False),
        (" to cut manual regression effort by ", False, False),
        ("50%", True, False),
        (".", False, False)
    ])

    # --- EDUCATION ---
    add_section_heading("Education")
    add_heading_row("Anna University \u2013 SIT, Tamil Nadu", "2007 \u2013 2011", is_bold=True, space_before=4)
    p_edu = doc.add_paragraph()
    p_edu.paragraph_format.space_before = Pt(0)
    p_edu.paragraph_format.space_after = Pt(2)
    p_edu.paragraph_format.left_indent = Inches(0.15)
    p_edu.add_run("Bachelor of Information Technology").font.italic = True

    # --- CERTIFICATIONS ---
    add_section_heading("Certifications")
    add_heading_row("Google Cloud Professional Data Engineer", "", is_bold=True, space_before=4)

    # --- AI WORKSPACE & TOOLING ---
    add_section_heading("AI Workspace & Tooling")
    add_heading_row("Preferred Developer Workspace & Agentic Tooling", "", space_before=4)
    add_bullet([
        ("Leverage ", False, False),
        ("Paseo ADE", True, False),
        (" for ACP agent orchestration (Codex, Claude Code) alongside Kiro-CLI, and OpenCode.", False, False)
    ])
    add_bullet([
        ("Configure agentic skills (such as superpowers) and custom MCP servers to automate coding and workspace workflows.", False, False)
    ])

    # Save to standard filename
    output_path = "Chokkar_Gurusamy_Resume.docx"
    doc.save(output_path)
    print(f"Successfully generated docx at: {os.path.abspath(output_path)}")

if __name__ == "__main__":
    create_resume()
